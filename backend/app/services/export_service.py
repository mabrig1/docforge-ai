"""
Export Service v2 — Generates .docx and .pdf with improved formatting fidelity.
"""

import re, io, os, subprocess, tempfile
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _hex_rgb(h: str) -> tuple:
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _sz(s) -> float:
    m = re.search(r"[\d.]+", str(s))
    return float(m.group()) if m else 12.0


def _align(s: str):
    return {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT, "justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}.get(s.lower(), WD_ALIGN_PARAGRAPH.JUSTIFY)


def _apply_run(run, font_name, size, bold=False, italic=False, color=None):
    """Apply formatting to a single run consistently."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        try:
            r, g, b = _hex_rgb(color)
            run.font.color.rgb = RGBColor(r, g, b)
        except (ValueError, IndexError):
            pass


def create_docx(structured_text: str, rules: dict) -> bytes:
    doc = Document()

    font = rules.get("font", "Times New Roman")
    bsz = _sz(rules.get("body_size", "12pt"))
    lsp = float(rules.get("line_spacing", "1.5"))
    psa = _sz(rules.get("paragraph_spacing_after", "8pt"))
    al = _align(rules.get("alignment", "justified"))
    margin = _sz(rules.get("margins", "1 inch"))
    h1 = rules.get("heading_1", {})
    h2 = rules.get("heading_2", {})
    h3 = rules.get("heading_3", {})
    indent_str = rules.get("first_line_indent", "none")
    indent = Inches(0.5) if indent_str and indent_str != "none" else None

    # Page margins
    for sec in doc.sections:
        sec.top_margin = Inches(margin)
        sec.bottom_margin = Inches(margin)
        sec.left_margin = Inches(margin)
        sec.right_margin = Inches(margin)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = font
    style.font.size = Pt(bsz)
    style.paragraph_format.line_spacing = lsp

    for line in structured_text.split("\n"):
        line = line.strip()
        if not line:
            continue

        tag_match = re.match(r"^\[([A-Z0-9]+)\]\s*(.*)", line, re.DOTALL)
        tag = tag_match.group(1) if tag_match else None
        text = tag_match.group(2).strip() if tag_match else line

        if tag == "TITLE":
            p = doc.add_paragraph()
            p.alignment = _align(h1.get("alignment", "center"))
            _apply_run(p.add_run(text), font, _sz(h1.get("size", "16pt")) + 4, bold=True, color=h1.get("color"))
            p.paragraph_format.space_after = Pt(24)
            p.paragraph_format.space_before = Pt(0)

        elif tag == "H1":
            p = doc.add_paragraph()
            p.alignment = _align(h1.get("alignment", "left"))
            caps = h1.get("caps", False)
            _apply_run(p.add_run(text.upper() if caps else text), font, _sz(h1.get("size", "16pt")),
                       bold=h1.get("bold", True), italic=h1.get("italic", False), color=h1.get("color"))
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)

        elif tag == "H2":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _apply_run(p.add_run(text), font, _sz(h2.get("size", "14pt")),
                       bold=h2.get("bold", True), italic=h2.get("italic", False), color=h2.get("color"))
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)

        elif tag == "H3":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _apply_run(p.add_run(text), font, _sz(h3.get("size", str(bsz) + "pt")),
                       bold=h3.get("bold", True), italic=h3.get("italic", True), color=h3.get("color"))
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)

        elif tag == "REFERENCES":
            doc.add_page_break()
            p = doc.add_paragraph()
            p.alignment = _align(h1.get("alignment", "center"))
            _apply_run(p.add_run(text or "References"), font, _sz(h1.get("size", "16pt")), bold=True, color=h1.get("color"))
            p.paragraph_format.space_after = Pt(16)

        elif tag == "REF":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _apply_run(p.add_run(text), font, bsz)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.line_spacing = lsp
            p.paragraph_format.space_after = Pt(4)

        else:  # PARAGRAPH or untagged
            p = doc.add_paragraph()
            p.alignment = al
            _apply_run(p.add_run(text), font, bsz)
            p.paragraph_format.line_spacing = lsp
            p.paragraph_format.space_after = Pt(psa)
            if indent:
                p.paragraph_format.first_line_indent = indent

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def create_pdf(structured_text: str, rules: dict) -> bytes:
    docx_bytes = create_docx(structured_text, rules)
    try:
        with tempfile.TemporaryDirectory() as tmp:
            dp = os.path.join(tmp, "document.docx")
            with open(dp, "wb") as f:
                f.write(docx_bytes)
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", dp, "--outdir", tmp],
                check=True, timeout=30, capture_output=True,
            )
            with open(os.path.join(tmp, "document.pdf"), "rb") as f:
                return f.read()
    except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
        return _pdf_reportlab(structured_text, rules)


def _pdf_reportlab(structured_text: str, rules: dict) -> bytes:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, PageBreak
    except ImportError as exc:
        raise RuntimeError("reportlab is not installed; PDF export is unavailable.") from exc

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=inch, bottomMargin=inch, leftMargin=inch, rightMargin=inch)
    bsz = _sz(rules.get("body_size", "12pt"))
    lh = bsz * float(rules.get("line_spacing", "1.5"))
    ss = getSampleStyleSheet()

    body = ParagraphStyle("Body", parent=ss["Normal"], fontName="Times-Roman", fontSize=bsz, leading=lh, spaceAfter=8)
    title = ParagraphStyle("T", parent=ss["Title"], fontName="Times-Bold", fontSize=bsz + 8, spaceAfter=24, alignment=1)
    h1s = ParagraphStyle("H1", parent=ss["Heading1"], fontName="Times-Bold", fontSize=_sz(rules.get("heading_1", {}).get("size", "16pt")), spaceBefore=18, spaceAfter=8)
    h2s = ParagraphStyle("H2", parent=ss["Heading2"], fontName="Times-Bold", fontSize=_sz(rules.get("heading_2", {}).get("size", "14pt")), spaceBefore=14, spaceAfter=6)
    h3s = ParagraphStyle("H3", parent=ss["Heading3"], fontName="Times-BoldItalic", fontSize=_sz(rules.get("heading_3", {}).get("size", str(bsz) + "pt")), spaceBefore=10, spaceAfter=4)
    refs = ParagraphStyle("Ref", parent=body, leftIndent=36, firstLineIndent=-36, spaceAfter=4)

    story = []
    for line in structured_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        m = re.match(r"^\[([A-Z0-9]+)\]\s*(.*)", line, re.DOTALL)
        tag = m.group(1) if m else None
        txt = m.group(2).strip() if m else line

        if tag == "TITLE": story.append(Paragraph(txt, title))
        elif tag == "H1": story.append(Paragraph(txt, h1s))
        elif tag == "H2": story.append(Paragraph(txt, h2s))
        elif tag == "H3": story.append(Paragraph(txt, h3s))
        elif tag == "REFERENCES": story += [PageBreak(), Paragraph(txt or "References", h1s)]
        elif tag == "REF": story.append(Paragraph(txt, refs))
        else: story.append(Paragraph(txt, body))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()
